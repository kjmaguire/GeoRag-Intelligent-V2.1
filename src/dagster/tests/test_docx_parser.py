"""Tests for docx_parser.py added in Sprint 3.

Covers:
  - parse_docx_report on real minimal .docx files (created via python-docx).
  - Resource table extraction: no-trigger vs trigger-containing table.
  - parse_doc_or_docx_report routing: .docx passthrough.
  - parse_doc_or_docx_report with .doc: libreoffice_unavailable, timeout, failed.
  - Unsupported extension raises ValueError.

Run with:  pytest tests/test_docx_parser.py -v
"""

from __future__ import annotations

import subprocess
from pathlib import Path
from unittest.mock import patch

import pytest

import docx as python_docx  # python-docx 1.2.0

from georag_dagster.parsers.docx_parser import (
    DocxParseResult,
    _sha256_of_file,
    parse_doc_or_docx_report,
    parse_docx_report,
)


# ---------------------------------------------------------------------------
# Fixtures: minimal .docx files built with python-docx at test time
# ---------------------------------------------------------------------------

@pytest.fixture()
def simple_docx(tmp_path: Path) -> Path:
    """A minimal .docx with a non-geological table — no trigger phrases,
    no column-header tokens — so resource_tables must be empty.
    """
    doc = python_docx.Document()
    doc.add_heading("Sample Technical Report", level=1)
    doc.add_paragraph("Prepared by Example Corp.")
    doc.add_paragraph("The project is located in Saskatchewan.")
    # Non-geological table: no trigger words, no column-header tokens.
    table = doc.add_table(rows=3, cols=3)
    for i, h in enumerate(["Name", "Date", "Reference"]):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text = "Smith, J."
    table.rows[1].cells[1].text = "2024-01-15"
    table.rows[1].cells[2].text = "Section 1"
    table.rows[2].cells[0].text = "Jones, A."
    table.rows[2].cells[1].text = "2024-03-20"
    table.rows[2].cells[2].text = "Section 2"
    p = tmp_path / "test.docx"
    doc.save(str(p))
    return p


@pytest.fixture()
def resource_docx(tmp_path: Path) -> Path:
    """A .docx whose TABLE CELLS contain the 'mineral resource' trigger phrase.

    _extract_tables_from_docx concatenates cell text to check for triggers;
    the trigger must appear in a cell (not just a paragraph) for trigger_phrase
    to be set on the result.
    """
    doc = python_docx.Document()
    doc.add_heading("NI 43-101 Technical Report", level=1)
    doc.add_paragraph("The project is located in Saskatchewan.")
    # Row 0 cell contains the trigger phrase so all_text check matches.
    table = doc.add_table(rows=4, cols=3)
    table.rows[0].cells[0].text = "Mineral Resource Estimate"
    table.rows[0].cells[1].text = ""
    table.rows[0].cells[2].text = ""
    for i, h in enumerate(["Category", "Tonnes", "Grade (g/t)"]):
        table.rows[1].cells[i].text = h
    table.rows[2].cells[0].text = "Measured"
    table.rows[2].cells[1].text = "5000000"
    table.rows[2].cells[2].text = "1.2"
    table.rows[3].cells[0].text = "Indicated"
    table.rows[3].cells[1].text = "3000000"
    table.rows[3].cells[2].text = "0.9"
    p = tmp_path / "resource.docx"
    doc.save(str(p))
    return p


@pytest.fixture()
def stub_doc(tmp_path: Path) -> Path:
    """A stub .doc file (OLE2 magic bytes — LibreOffice absent so not parsed)."""
    p = tmp_path / "legacy.doc"
    p.write_bytes(b"\xD0\xCF\x11\xE0" + b"\x00" * 100)
    return p


# ---------------------------------------------------------------------------
# parse_docx_report — basic shape
# ---------------------------------------------------------------------------

class TestParseDocxReportShape:
    def test_full_text_contains_paragraph_text(self, simple_docx: Path):
        result = parse_docx_report(simple_docx)
        assert "Saskatchewan" in result.full_text
        assert "Example Corp" in result.full_text

    def test_title_extracted_from_heading(self, simple_docx: Path):
        result = parse_docx_report(simple_docx)
        assert result.title == "Sample Technical Report"

    def test_provenance_sha256_is_64_char_hex(self, simple_docx: Path):
        result = parse_docx_report(simple_docx)
        sha = result.provenance.get("source_file_sha256", "")
        assert len(sha) == 64
        assert all(c in "0123456789abcdef" for c in sha)

    def test_provenance_parser_name(self, simple_docx: Path):
        result = parse_docx_report(simple_docx)
        assert result.provenance.get("parser_name") == "docx_parser"

    def test_parser_used_is_python_docx(self, simple_docx: Path):
        result = parse_docx_report(simple_docx)
        assert result.parser_used == "python-docx"

    def test_is_scanned_false(self, simple_docx: Path):
        result = parse_docx_report(simple_docx)
        assert result.is_scanned is False

    def test_warnings_is_list(self, simple_docx: Path):
        result = parse_docx_report(simple_docx)
        assert isinstance(result.warnings, list)

    def test_resource_tables_empty_without_trigger(self, simple_docx: Path):
        """The simple_docx fixture has no trigger phrases and no geological
        column-header tokens → _extract_tables_from_docx returns empty list."""
        result = parse_docx_report(simple_docx)
        assert result.resource_tables == [], (
            "Expected empty resource_tables when no trigger phrase and no geological "
            f"column-header tokens appear in table cells; got: {result.resource_tables}"
        )

    def test_result_is_docx_parse_result(self, simple_docx: Path):
        result = parse_docx_report(simple_docx)
        assert isinstance(result, DocxParseResult)

    def test_file_not_found_raises(self, tmp_path: Path):
        with pytest.raises(FileNotFoundError):
            parse_docx_report(tmp_path / "ghost.docx")


# ---------------------------------------------------------------------------
# parse_docx_report — resource table extraction
# ---------------------------------------------------------------------------

class TestParseDocxReportResourceTables:
    def test_resource_table_extracted_when_trigger_present(self, resource_docx: Path):
        """Docx with 'Mineral Resource Estimate' trigger in cell → resource_tables non-empty."""
        result = parse_docx_report(resource_docx)
        assert len(result.resource_tables) >= 1

    def test_resource_table_extraction_method_is_docx_native(self, resource_docx: Path):
        result = parse_docx_report(resource_docx)
        for tbl in result.resource_tables:
            assert tbl["extraction_method"] == "docx_native"

    def test_resource_table_confidence_is_float_in_range(self, resource_docx: Path):
        result = parse_docx_report(resource_docx)
        for tbl in result.resource_tables:
            assert isinstance(tbl["confidence"], float)
            assert 0.0 <= tbl["confidence"] <= 1.0

    def test_resource_table_has_required_keys(self, resource_docx: Path):
        result = parse_docx_report(resource_docx)
        required_keys = {
            "page", "table_index_on_page", "trigger_phrase",
            "header", "rows", "extraction_method", "confidence",
        }
        for tbl in result.resource_tables:
            assert required_keys.issubset(tbl.keys())

    def test_resource_table_trigger_phrase_set(self, resource_docx: Path):
        """trigger_phrase must be non-None when the trigger appears in table cells."""
        result = parse_docx_report(resource_docx)
        assert len(result.resource_tables) >= 1
        tbl = result.resource_tables[0]
        assert tbl["trigger_phrase"] is not None, (
            "trigger_phrase should be set when 'Mineral Resource Estimate' appears in cell text"
        )


# ---------------------------------------------------------------------------
# parse_doc_or_docx_report — routing and failure paths
# ---------------------------------------------------------------------------

class TestParseDocOrDocxReportRouting:
    def test_docx_extension_routes_to_parse_docx_report(self, simple_docx: Path):
        """.docx path goes directly to parse_docx_report — same output."""
        result = parse_doc_or_docx_report(simple_docx)
        assert isinstance(result, DocxParseResult)
        assert result.parser_used == "python-docx"
        assert "Saskatchewan" in result.full_text

    def test_unsupported_extension_raises_value_error(self, tmp_path: Path):
        bad = tmp_path / "report.pdf"
        bad.write_bytes(b"%PDF-1.4\n")
        with pytest.raises(ValueError, match="unsupported extension"):
            parse_doc_or_docx_report(bad)

    def test_txt_extension_raises_value_error(self, tmp_path: Path):
        bad = tmp_path / "report.txt"
        bad.write_text("hello")
        with pytest.raises(ValueError):
            parse_doc_or_docx_report(bad)


class TestParseDocOrDocxReportLibreofficeUnavailable:
    def test_no_soffice_returns_structured_failure(self, stub_doc: Path):
        """When shutil.which returns None for soffice and libreoffice,
        result must NOT raise and must carry 'libreoffice_unavailable' warning."""
        with patch("shutil.which", return_value=None):
            result = parse_doc_or_docx_report(stub_doc)

        assert isinstance(result, DocxParseResult)
        assert result.full_text == ""
        assert result.parser_used == "libreoffice_conversion_skipped"
        codes = [w.get("code") for w in result.warnings if isinstance(w, dict)]
        assert "libreoffice_unavailable" in codes

    def test_no_soffice_result_does_not_raise(self, stub_doc: Path):
        with patch("shutil.which", return_value=None):
            try:
                parse_doc_or_docx_report(stub_doc)
                did_raise = False
            except Exception:
                did_raise = True
        assert not did_raise


class TestParseDocOrDocxReportLibreofficeTimeout:
    def test_timeout_returns_libreoffice_conversion_timeout_warning(
        self, stub_doc: Path
    ):
        """subprocess.run raising TimeoutExpired → structured failure, no raise."""
        with patch("shutil.which", return_value="/usr/bin/soffice"):
            with patch(
                "subprocess.run",
                side_effect=subprocess.TimeoutExpired(cmd="soffice", timeout=120),
            ):
                result = parse_doc_or_docx_report(stub_doc)

        assert isinstance(result, DocxParseResult)
        codes = [w.get("code") for w in result.warnings if isinstance(w, dict)]
        assert "libreoffice_conversion_timeout" in codes

    def test_timeout_does_not_raise(self, stub_doc: Path):
        with patch("shutil.which", return_value="/usr/bin/soffice"):
            with patch(
                "subprocess.run",
                side_effect=subprocess.TimeoutExpired(cmd="soffice", timeout=120),
            ):
                try:
                    parse_doc_or_docx_report(stub_doc)
                    did_raise = False
                except Exception:
                    did_raise = True
        assert not did_raise


class TestParseDocOrDocxReportLibreofficeCalledProcessError:
    def test_called_process_error_returns_libreoffice_conversion_failed_warning(
        self, stub_doc: Path
    ):
        """subprocess.CalledProcessError → structured failure, no raise."""
        with patch("shutil.which", return_value="/usr/bin/soffice"):
            with patch(
                "subprocess.run",
                side_effect=subprocess.CalledProcessError(
                    returncode=1, cmd="soffice", stderr=b"conversion error"
                ),
            ):
                result = parse_doc_or_docx_report(stub_doc)

        assert isinstance(result, DocxParseResult)
        codes = [w.get("code") for w in result.warnings if isinstance(w, dict)]
        assert "libreoffice_conversion_failed" in codes

    def test_called_process_error_does_not_raise(self, stub_doc: Path):
        with patch("shutil.which", return_value="/usr/bin/soffice"):
            with patch(
                "subprocess.run",
                side_effect=subprocess.CalledProcessError(
                    returncode=1, cmd="soffice"
                ),
            ):
                try:
                    parse_doc_or_docx_report(stub_doc)
                    did_raise = False
                except Exception:
                    did_raise = True
        assert not did_raise


# ---------------------------------------------------------------------------
# _sha256_of_file
# ---------------------------------------------------------------------------

class TestSha256OfFile:
    def test_returns_64_char_hex(self, simple_docx: Path):
        digest = _sha256_of_file(simple_docx)
        assert len(digest) == 64
        assert all(c in "0123456789abcdef" for c in digest)

    def test_deterministic(self, simple_docx: Path):
        assert _sha256_of_file(simple_docx) == _sha256_of_file(simple_docx)

    def test_different_files_different_digest(self, tmp_path: Path):
        a = tmp_path / "a.txt"
        b = tmp_path / "b.txt"
        a.write_bytes(b"content_a")
        b.write_bytes(b"content_b")
        assert _sha256_of_file(a) != _sha256_of_file(b)
