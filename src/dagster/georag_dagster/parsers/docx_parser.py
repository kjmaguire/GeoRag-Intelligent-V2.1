"""NI 43-101 DOCX/DOC Report Parser — Bronze → Silver ingestion for Word documents.

Handles both modern .docx files (via python-docx) and legacy .doc files (via
LibreOffice subprocess conversion). Outputs a DocxParseResult whose shape
mirrors ReportParseResult so downstream consumers (silver_reports asset) can
treat both without branching.

LibreOffice (soffice) is an optional runtime dependency — if the binary is
absent the .doc path returns a clean structured failure rather than raising,
so the pipeline degrades gracefully. Installing libreoffice-writer is a
devops-engineer Dockerfile task tracked separately.

License: python-docx is MIT-licensed (confirmed pip show python-docx).
"""

from __future__ import annotations

import hashlib
import logging
import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

from georag_dagster.parsers.pdf_report import (
    ReportSection,
    _COLUMN_HEADER_TOKENS,
    _RESOURCE_TABLE_TRIGGERS,
    _detect_page_language,
    _extract_authors,
    _extract_commodity,
    _extract_company,
    _extract_filing_date,
    _extract_project_name,
    _extract_region,
    _split_into_sections,
    _table_confidence,
)

logger = logging.getLogger(__name__)

PARSER_VERSION = "1.0.0"


# ---------------------------------------------------------------------------
# Result dataclass — mirrors ReportParseResult for downstream compat
# ---------------------------------------------------------------------------

@dataclass
class DocxParseResult:
    """Complete result of parsing a NI 43-101 Word document (DOCX or DOC)."""

    full_text: str
    title: Optional[str]
    sections: list[ReportSection]
    company: Optional[str]
    filing_date: Optional[str]      # ISO 8601 string: YYYY-MM-DD
    commodity: Optional[str]
    authors: list[str]
    project_name: Optional[str]
    region: Optional[str]
    warnings: list[dict] = field(default_factory=list)
    provenance: dict[str, Any] = field(default_factory=dict)
    resource_tables: list[dict] = field(default_factory=list)
    page_languages: list[str] = field(default_factory=list)  # best-effort; whole-doc one entry
    parse_quality_pct: float = 0.0
    extraction_confidence: float = 1.0
    parser_used: str = "python-docx"
    is_scanned: bool = False


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _sha256_of_file(path: Path) -> str:
    """Return the SHA-256 hex digest of the file at *path*."""
    h = hashlib.sha256()
    with open(path, "rb") as fh:
        for chunk in iter(lambda: fh.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def _extract_tables_from_docx(doc) -> list[dict]:
    """Extract tables from a python-docx Document, applying resource-table heuristics.

    For each table in doc.tables, concatenates all cell text and checks whether
    any resource-table trigger phrase is present. If so (or when more than half
    the column headers match known tokens), the table is included as a
    structured entry.

    Returns a list of table dicts compatible with the format produced by
    _extract_resource_tables() in pdf_report.py.
    """
    results: list[dict] = []

    for tbl_idx, table in enumerate(doc.tables):
        # Build a list-of-lists (str) from the table
        raw_rows: list[list[str]] = []
        for row in table.rows:
            raw_rows.append([cell.text.strip() for cell in row.cells])

        if not raw_rows:
            continue

        # Flatten table text to check for resource triggers
        all_text = " ".join(cell for row in raw_rows for cell in row).lower()
        matched_trigger: Optional[str] = None
        for trigger in _RESOURCE_TABLE_TRIGGERS:
            if trigger in all_text:
                matched_trigger = trigger
                break

        # Use row 0 as header candidate; score it
        header_row = raw_rows[0]
        header_score = sum(
            1 for cell in header_row
            if any(tok in cell.lower() for tok in _COLUMN_HEADER_TOKENS)
        )
        header_fraction = header_score / max(len(header_row), 1)

        # Include the table if a trigger matched OR if >30% of header cells
        # match known column-header tokens (catches unlabelled resource tables)
        if matched_trigger is None and header_fraction < 0.3:
            continue

        # Clean up header: replace empty cells with positional fallback
        header = [cell if cell else f"col_{i}" for i, cell in enumerate(header_row)]
        data_rows = raw_rows[1:]

        confidence = _table_confidence(header, data_rows)
        # Docx native extraction is higher-fidelity than PDF heuristics; bump floor
        confidence = max(confidence, 0.9 if matched_trigger else 0.5)

        results.append({
            "page": None,                       # DOCX has no page structure
            "table_index_on_page": tbl_idx,
            "trigger_phrase": matched_trigger,
            "header": header,
            "rows": data_rows,
            "extraction_method": "docx_native",
            "confidence": round(min(1.0, confidence), 4),
        })

    return results


def _parse_quality_from_sections(sections: list[ReportSection]) -> float:
    """Compute parse quality as fraction of 17 expected NI 43-101 sections found."""
    from georag_dagster.parsers.pdf_report import NI43_BASELINE_SECTIONS  # noqa: PLC0415
    numbered = [s for s in sections if s.section_number is not None]
    return round(len(numbered) / NI43_BASELINE_SECTIONS, 4)


# ---------------------------------------------------------------------------
# Public: DOCX parser
# ---------------------------------------------------------------------------

def parse_docx_report(path: str | Path) -> DocxParseResult:
    """Parse a NI 43-101 DOCX technical report and return a :class:`DocxParseResult`.

    Parameters
    ----------
    path:
        Absolute path to a .docx file on disk.

    Returns
    -------
    DocxParseResult
        Extracted metadata, sections, resource tables, and parse quality.
    """
    import docx  # python-docx; noqa: PLC0415

    p = Path(path)
    if not p.is_file():
        raise FileNotFoundError(f"parse_docx_report: file not found at '{p}'")

    sha256 = _sha256_of_file(p)
    provenance: dict[str, Any] = {
        "source_file": str(p),
        "source_file_sha256": sha256,
        "parser_name": "docx_parser",
        "parser_version": PARSER_VERSION,
    }

    warnings: list[dict] = []

    try:
        doc = docx.Document(str(p))
    except Exception as exc:
        warnings.append({
            "code": "docx_open_failed",
            "message": str(exc),
        })
        return DocxParseResult(
            full_text="",
            title=None,
            sections=[],
            company=None,
            filing_date=None,
            commodity=None,
            authors=[],
            project_name=None,
            region=None,
            warnings=warnings,
            provenance=provenance,
            parse_quality_pct=0.0,
            parser_used="python-docx",
        )

    # --- Full text: join non-empty paragraph texts ---
    full_text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

    # --- Section splitting: headings drive structure ---
    # Build a heading-aware text block where heading paragraphs are formatted
    # to match SECTION_HEADING_RE so _split_into_sections can reuse the same
    # regex. e.g. a Heading 1 para "14. MINERAL RESOURCE ESTIMATES" passes
    # directly; plain headings get normalised here if they start with a digit.
    heading_aware_lines: list[str] = []
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            # Headings are already valid section markers — preserve at line start
            heading_aware_lines.append(para.text.strip())
        else:
            heading_aware_lines.append(para.text.strip())
    heading_aware_text = "\n".join(heading_aware_lines)

    sections = _split_into_sections(heading_aware_text)

    # --- Metadata extraction from lead text ---
    lead_text = full_text[:2000]
    title_candidate = ""
    for para in doc.paragraphs:
        if para.text.strip():
            title_candidate = para.text.strip()[:200]
            break

    authors = _extract_authors(lead_text)
    company = _extract_company(lead_text)
    filing_date = _extract_filing_date(lead_text)
    commodity = _extract_commodity(lead_text) or _extract_commodity(full_text[:5000])
    project_name = _extract_project_name(lead_text, title_candidate or None)
    region = _extract_region(lead_text) or _extract_region(full_text[:5000])

    # --- Language detection (whole-doc; one entry) ---
    lang = _detect_page_language(full_text)
    page_languages = [lang]

    # --- Resource table extraction ---
    resource_tables: list[dict] = []
    try:
        resource_tables = _extract_tables_from_docx(doc)
        logger.info(
            "docx_parser: resource table extraction found %d table(s) in '%s'",
            len(resource_tables),
            p.name,
        )
    except Exception as tbl_exc:
        warnings.append({
            "code": "resource_table_extraction_failed",
            "message": str(tbl_exc),
        })
        logger.warning(
            "docx_parser: resource table extraction failed for '%s': %s",
            p.name,
            tbl_exc,
        )

    parse_quality_pct = _parse_quality_from_sections(sections)
    metadata_fields = sum(
        1 for v in [title_candidate, authors, company, filing_date, commodity, project_name, region]
        if v
    )
    extraction_confidence = min(1.0, (
        parse_quality_pct * 0.5
        + min(1.0, len(full_text) / 10000) * 0.3
        + (metadata_fields / 7) * 0.2
    ))

    logger.info(
        "docx_parser: parse complete — sections=%d, quality=%.1f%%, "
        "title='%s', commodity=%s",
        len(sections),
        parse_quality_pct * 100,
        (title_candidate or "")[:60],
        commodity,
    )

    return DocxParseResult(
        full_text=full_text,
        title=title_candidate or None,
        sections=sections,
        company=company,
        filing_date=filing_date,
        commodity=commodity,
        authors=authors,
        project_name=project_name,
        region=region,
        warnings=warnings,
        provenance=provenance,
        resource_tables=resource_tables,
        page_languages=page_languages,
        parse_quality_pct=parse_quality_pct,
        extraction_confidence=extraction_confidence,
        parser_used="python-docx",
        is_scanned=False,
    )


# ---------------------------------------------------------------------------
# Public: DOC fallback via LibreOffice subprocess
# ---------------------------------------------------------------------------

def parse_doc_or_docx_report(path: str | Path) -> DocxParseResult:
    """Entry point that routes .doc files through LibreOffice first, then reuses the .docx path.

    .docx files are parsed directly. .doc files are converted via the
    LibreOffice headless binary (soffice / libreoffice) and then parsed as
    .docx. If the binary is absent, a clean structured failure is returned
    rather than raising — the caller must check result.warnings for the
    "libreoffice_unavailable" code.

    Parameters
    ----------
    path:
        Absolute path to a .doc or .docx file on disk.

    Returns
    -------
    DocxParseResult
        Parsed report or a failure stub with a structured warning.
    """
    p = Path(path)

    if p.suffix.lower() == ".docx":
        return parse_docx_report(p)

    if p.suffix.lower() != ".doc":
        raise ValueError(f"parse_doc_or_docx_report: unsupported extension '{p.suffix}'")

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice is None:
        # Clean structured failure — does NOT raise.
        logger.warning(
            "docx_parser: soffice/libreoffice not found — cannot convert '%s'", p.name
        )
        return DocxParseResult(
            full_text="",
            title=None,
            sections=[],
            company=None,
            filing_date=None,
            commodity=None,
            authors=[],
            project_name=None,
            region=None,
            warnings=[{
                "code": "libreoffice_unavailable",
                "message": (
                    ".doc files require LibreOffice (soffice binary). "
                    "Install libreoffice-writer."
                ),
                "context": {"path": str(p)},
            }],
            provenance={
                "parser_name": "docx_parser",
                "parser_version": PARSER_VERSION,
                "source_file": str(p),
                "source_file_sha256": _sha256_of_file(p),
            },
            parse_quality_pct=0.0,
            parser_used="libreoffice_conversion_skipped",
        )

    # Convert in a temp dir so we don't pollute the source.
    with tempfile.TemporaryDirectory() as tmp:
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "docx", "--outdir", tmp, str(p)],
                check=True,
                capture_output=True,
                timeout=120,
            )
        except subprocess.TimeoutExpired:
            logger.error(
                "docx_parser: LibreOffice conversion timed out after 120s for '%s'", p.name
            )
            return DocxParseResult(
                full_text="",
                title=None,
                sections=[],
                company=None,
                filing_date=None,
                commodity=None,
                authors=[],
                project_name=None,
                region=None,
                warnings=[{
                    "code": "libreoffice_conversion_timeout",
                    "message": "LibreOffice conversion exceeded 120 s timeout.",
                    "context": {"path": str(p)},
                }],
                provenance={
                    "parser_name": "docx_parser",
                    "parser_version": PARSER_VERSION,
                    "source_file": str(p),
                    "source_file_sha256": _sha256_of_file(p),
                },
                parse_quality_pct=0.0,
                parser_used="libreoffice_conversion_timeout",
            )
        except subprocess.CalledProcessError as cpe:
            stderr_text = cpe.stderr.decode(errors="replace") if cpe.stderr else ""
            logger.error(
                "docx_parser: LibreOffice conversion failed for '%s' (exit %d): %s",
                p.name,
                cpe.returncode,
                stderr_text[:500],
            )
            return DocxParseResult(
                full_text="",
                title=None,
                sections=[],
                company=None,
                filing_date=None,
                commodity=None,
                authors=[],
                project_name=None,
                region=None,
                warnings=[{
                    "code": "libreoffice_conversion_failed",
                    "message": f"soffice exited {cpe.returncode}",
                    "context": {"stderr": stderr_text[:500], "path": str(p)},
                }],
                provenance={
                    "parser_name": "docx_parser",
                    "parser_version": PARSER_VERSION,
                    "source_file": str(p),
                    "source_file_sha256": _sha256_of_file(p),
                },
                parse_quality_pct=0.0,
                parser_used="libreoffice_conversion_failed",
            )

        converted = Path(tmp) / (p.stem + ".docx")
        if not converted.exists():
            # LibreOffice sometimes conjures different filenames on odd inputs
            candidates = list(Path(tmp).glob("*.docx"))
            if not candidates:
                raise RuntimeError(
                    f"docx_parser: LibreOffice conversion produced no .docx for '{p}'"
                )
            converted = candidates[0]

        result = parse_docx_report(converted)

    # Overwrite provenance with the original .doc metadata, not the temp .docx
    result.provenance["source_file"] = str(p)
    result.provenance["source_file_sha256"] = _sha256_of_file(p)
    result.parser_used = "libreoffice+python-docx"
    return result
